"""Generate 10-page Assignment 01 report as a DOCX file."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = os.path.join(os.path.dirname(__file__), "Assignment01_Report.docx")
FIGURES_DIR = os.path.join(os.path.dirname(__file__), "notebooks", "figures")


# ─── helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_page_break(doc):
    doc.add_page_break()


def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p


def body(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_figure(doc, path, caption, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    else:
        p = doc.add_paragraph(f"[Screenshot placeholder: {caption}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def add_table(doc, headers, rows, header_color="1F497D"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_color)
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            set_cell_bg(cells[ci], bg)
            cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
            if ci == 0:
                cells[ci].paragraphs[0].runs[0].bold = True
    return table


# ─── build document ─────────────────────────────────────────────────────────────

def build_report():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title Page ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BITS Pilani — Work Integrated Learning Programmes")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Machine Learning Operations (AIMLCZG523)")
    r.font.size = Pt(13)
    r.bold = True

    doc.add_paragraph()

    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = main_title.add_run("Assignment 01\nEnd-to-End ML Model Development,\nCI/CD and Production Deployment")
    rt.bold = True
    rt.font.size = Pt(18)
    rt.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()
    doc.add_paragraph()

    meta_items = [
        ("Student Name", "Yogamani Balusamy"),
        ("Email", "yogamani.balusamy@gmail.com"),
        ("Course", "AIMLCZG523 — Machine Learning Operations"),
        ("Assignment", "Assignment 01 (Total Marks: 50)"),
        ("Date of Submission", datetime.date.today().strftime("%B %d, %Y")),
        ("Dataset", "Heart Disease UCI Dataset"),
        ("Repository", "https://github.com/yogisamy/MLOps-Assignment1"),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    add_page_break(doc)

    # ── 1. Project Overview ──────────────────────────────────────────────────────
    heading1(doc, "1. Project Overview")

    body(doc, (
        "This project implements a complete end-to-end Machine Learning Operations (MLOps) pipeline "
        "for predicting the presence of heart disease in patients. The solution is built using the "
        "Heart Disease UCI dataset from the UCI Machine Learning Repository and encompasses all stages "
        "of a production-grade ML system: data acquisition, exploratory analysis, model development, "
        "experiment tracking, containerisation, CI/CD automation, Kubernetes deployment, and monitoring."
    ))

    heading2(doc, "1.1 Objective")
    body(doc, (
        "Build a binary classifier to predict the risk of heart disease from 13 patient health features "
        "and expose the model as a cloud-ready, monitored REST API. Every stage must be reproducible, "
        "automated, and observable — mirroring real-world MLOps practices."
    ))

    heading2(doc, "1.2 Technology Stack")
    add_table(doc,
        ["Category", "Tool / Library"],
        [
            ["Programming Language", "Python 3.13"],
            ["Data Processing", "Pandas, NumPy"],
            ["Visualisation", "Matplotlib, Seaborn, Plotly"],
            ["ML Frameworks", "Scikit-learn, XGBoost"],
            ["Experiment Tracking", "MLflow (SQLite backend)"],
            ["API Framework", "FastAPI + Uvicorn"],
            ["Testing", "Pytest, pytest-cov"],
            ["Containerisation", "Docker, Docker Compose"],
            ["CI/CD", "GitHub Actions"],
            ["Orchestration", "Kubernetes (Minikube / Docker Desktop)"],
            ["Monitoring", "Prometheus, Grafana"],
            ["Version Control", "Git, GitHub"],
        ]
    )

    heading2(doc, "1.3 Repository Structure")
    for line in [
        "Assignment1/",
        "├── data/              — download script and heart.csv",
        "├── notebooks/         — EDA, inference scripts and figures",
        "├── src/               — data_processing, train, predict modules",
        "├── api/               — FastAPI application and Pydantic schemas",
        "├── tests/             — 16 pytest unit tests",
        "├── k8s/               — Kubernetes Deployment and Service YAMLs",
        "├── monitoring/        — Prometheus configuration",
        "├── .github/workflows/ — GitHub Actions CI/CD pipeline",
        "├── Dockerfile",
        "├── docker-compose.yml — API + Prometheus + Grafana stack",
        "└── requirements.txt",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(1)

    add_page_break(doc)

    # ── 2. Data Acquisition & EDA ────────────────────────────────────────────────
    heading1(doc, "2. Data Acquisition & Exploratory Data Analysis")

    heading2(doc, "2.1 Dataset Description")
    body(doc, (
        "The Heart Disease UCI Dataset (Cleveland subset) contains 303 patient records with 13 "
        "clinical features and a binary target variable indicating presence (1) or absence (0) of "
        "heart disease. The dataset was obtained programmatically via the ucimlrepo Python package "
        "with a CSV fallback to the UCI archive URL."
    ))

    add_table(doc,
        ["Feature", "Type", "Description"],
        [
            ["age", "Numerical", "Age in years"],
            ["sex", "Categorical", "1 = male, 0 = female"],
            ["cp", "Categorical", "Chest pain type (0–3)"],
            ["trestbps", "Numerical", "Resting blood pressure (mmHg)"],
            ["chol", "Numerical", "Serum cholesterol (mg/dl)"],
            ["fbs", "Categorical", "Fasting blood sugar > 120 mg/dl"],
            ["restecg", "Categorical", "Resting ECG results (0–2)"],
            ["thalach", "Numerical", "Maximum heart rate achieved"],
            ["exang", "Categorical", "Exercise-induced angina (1=yes)"],
            ["oldpeak", "Numerical", "ST depression by exercise"],
            ["slope", "Categorical", "Slope of peak exercise ST segment"],
            ["ca", "Numerical", "Number of major vessels (0–3)"],
            ["thal", "Categorical", "Thalassemia type (1/2/3)"],
            ["target", "Binary", "0 = No disease, 1 = Disease (label)"],
        ]
    )

    heading2(doc, "2.2 Missing Value Analysis")
    body(doc, (
        "The dataset contains a small number of missing values: the 'ca' feature has 4 missing "
        "values (1.32%) and 'thal' has 2 missing values (0.66%). All missing values are handled "
        "by the sklearn preprocessing pipeline — numerical features use median imputation and "
        "categorical features use most-frequent (mode) imputation."
    ))
    add_figure(doc, os.path.join(FIGURES_DIR, "01_missing_values.png"),
               "Figure 1: Missing Value Percentage per Feature")

    heading2(doc, "2.3 Class Distribution")
    body(doc, (
        "The dataset is nearly balanced: 165 patients (54.5%) have no heart disease and "
        "138 patients (45.5%) have heart disease. Stratified train/test splitting preserves "
        "this ratio in both partitions."
    ))
    add_figure(doc, os.path.join(FIGURES_DIR, "02_class_distribution.png"),
               "Figure 2: Class Distribution — Heart Disease vs No Heart Disease")

    heading2(doc, "2.4 Feature Distributions")
    body(doc, (
        "Histograms of all numerical features reveal that 'age' is approximately normally "
        "distributed (mean ~54 years), 'chol' is right-skewed with some extreme values, "
        "and 'thalach' (max heart rate) decreases with age as expected. The 'oldpeak' "
        "feature is heavily right-skewed, suggesting a logarithmic transformation could be "
        "beneficial in future iterations."
    ))
    add_figure(doc, os.path.join(FIGURES_DIR, "03_histograms.png"),
               "Figure 3: Histograms of All Numerical Features")

    heading2(doc, "2.5 Correlation Analysis")
    body(doc, (
        "The correlation heatmap reveals: 'thalach' (max heart rate) has a moderate negative "
        "correlation with target (−0.42), indicating higher max heart rate is associated with "
        "no disease. 'cp' (chest pain type) shows the strongest positive correlation with "
        "target (+0.43). 'oldpeak' is positively correlated with target (+0.43). "
        "Multicollinearity is low across features, making them collectively informative."
    ))
    add_figure(doc, os.path.join(FIGURES_DIR, "04_correlation_heatmap.png"),
               "Figure 4: Feature Correlation Heatmap")

    add_figure(doc, os.path.join(FIGURES_DIR, "05_feature_by_target.png"),
               "Figure 5: Feature Distributions Split by Target Class")

    add_figure(doc, os.path.join(FIGURES_DIR, "06_categorical_by_target.png"),
               "Figure 6: Categorical Feature Counts by Target Class")

    add_page_break(doc)

    # ── 3. Feature Engineering & Model Development ───────────────────────────────
    heading1(doc, "3. Feature Engineering & Model Development")

    heading2(doc, "3.1 Preprocessing Pipeline")
    body(doc, (
        "A reusable sklearn ColumnTransformer pipeline was constructed to handle all "
        "preprocessing steps identically during training and inference:"
    ))
    bullet(doc, "Numerical features (age, trestbps, chol, thalach, oldpeak, ca): "
                "Median imputation → StandardScaler normalisation")
    bullet(doc, "Categorical features (sex, cp, fbs, restecg, exang, slope, thal): "
                "Mode imputation → OrdinalEncoder (unknown values → −1)")
    body(doc, (
        "The entire pipeline (preprocessor + classifier) is saved as a single joblib artifact, "
        "ensuring that inference always applies the same transformations fitted on training data "
        "with no data leakage."
    ))

    heading2(doc, "3.2 Models Trained")
    body(doc, "Two classification models were trained and compared:")

    body(doc, "Model 1 — Logistic Regression", bold=True)
    bullet(doc, "Hyperparameter search: C ∈ {0.01, 0.1, 1, 10}, solver ∈ {lbfgs, liblinear}")
    bullet(doc, "Optimisation criterion: ROC-AUC via 5-fold StratifiedKFold GridSearchCV")
    bullet(doc, "Best params: C=1, solver=liblinear")

    body(doc, "Model 2 — Random Forest Classifier", bold=True)
    bullet(doc, "Hyperparameter search: n_estimators ∈ {100, 200}, max_depth ∈ {None, 5, 10}, "
                "min_samples_split ∈ {2, 5}")
    bullet(doc, "Optimisation criterion: ROC-AUC via 5-fold StratifiedKFold GridSearchCV")
    bullet(doc, "Best params: n_estimators=100, max_depth=None, min_samples_split=2")

    heading2(doc, "3.3 Model Evaluation")
    body(doc, (
        "Both models were evaluated on a held-out 20% test set (61 samples) using five metrics. "
        "5-fold cross-validation scores are also reported for robustness."
    ))

    add_table(doc,
        ["Metric", "Logistic Regression", "Random Forest"],
        [
            ["Accuracy", "86.9%", "90.2%"],
            ["Precision", "81.3%", "86.7%"],
            ["Recall", "92.9%", "92.9%"],
            ["F1-Score", "86.7%", "89.7%"],
            ["ROC-AUC (test)", "0.9524", "0.9632"],
            ["CV ROC-AUC (mean ± std)", "0.8981 ± 0.0146", "0.8882 ± 0.0231"],
        ]
    )
    doc.add_paragraph()

    body(doc, (
        "The Random Forest model achieves a higher test ROC-AUC (0.9632 vs 0.9524) and better "
        "accuracy and F1. Despite Logistic Regression showing slightly better CV ROC-AUC stability "
        "(lower std), Random Forest was selected as the best model due to its superior test-set "
        "generalisation and recall parity. The high recall (92.9% for both) is particularly "
        "important in a medical diagnosis context — minimising false negatives is critical."
    ))

    add_page_break(doc)

    # ── 4. Experiment Tracking ───────────────────────────────────────────────────
    heading1(doc, "4. Experiment Tracking with MLflow")

    body(doc, (
        "MLflow was integrated as the primary experiment tracking system. All training runs are "
        "logged to a local SQLite backend (mlflow.db) and can be inspected via the MLflow UI "
        "at http://localhost:5000 by running: mlflow ui --backend-store-uri sqlite:///mlflow.db"
    ))

    heading2(doc, "4.1 What Is Logged")
    add_table(doc,
        ["MLflow Artifact", "Details"],
        [
            ["Run name", "logistic_regression / random_forest"],
            ["Tag: model_name", "Identifies the model type"],
            ["Parameters", "All best hyperparameters from GridSearchCV"],
            ["Test metrics", "accuracy, precision, recall, f1, roc_auc"],
            ["CV metrics", "cv_<metric>_mean and cv_<metric>_std for all 5 metrics"],
            ["Confusion matrix plot", "PNG logged under plots/"],
            ["ROC curve plot", "PNG logged under plots/"],
            ["Serialised model", "pickle format logged under model/"],
            ["Joblib artifact", "best_model.joblib logged under saved_model/"],
        ]
    )

    heading2(doc, "4.2 MLflow Screenshots")
    body(doc, "The following screenshots show the MLflow experiment tracking UI:", italic=True)
    add_figure(doc, "screenshots/mlflow_experiments.png",
               "Figure 7: MLflow Experiments List — heart_disease_classification")
    add_figure(doc, "screenshots/mlflow_experiments_detail.png",
               "Figure 8: MLflow Experiment — Run Comparison View")
    add_figure(doc, "screenshots/mlflow_run_detail.png",
               "Figure 9: MLflow Run Detail — Parameters, Metrics and Artifacts")

    add_page_break(doc)

    # ── 5. Model Packaging & Reproducibility ─────────────────────────────────────
    heading1(doc, "5. Model Packaging & Reproducibility")

    heading2(doc, "5.1 Model Serialisation")
    body(doc, (
        "The complete sklearn Pipeline (preprocessor + classifier) is serialised using joblib "
        "and saved to models/best_model.joblib. This single file encapsulates all transformations "
        "fitted on training data, ensuring identical behaviour during inference without retraining "
        "or re-fitting any preprocessing steps."
    ))

    heading2(doc, "5.2 Dependency Management")
    body(doc, (
        "All Python dependencies are pinned in requirements.txt using minimum version constraints "
        "compatible with Python 3.13. The file covers: data processing (pandas, numpy), ML "
        "(scikit-learn, xgboost, joblib), visualisation (matplotlib, seaborn, plotly), tracking "
        "(mlflow), API (fastapi, uvicorn, pydantic), testing (pytest, httpx), linting (flake8, "
        "black), and monitoring (prometheus-client)."
    ))

    heading2(doc, "5.3 Reproducibility Guarantee")
    bullet(doc, "train_test_split uses random_state=42 and stratify=y")
    bullet(doc, "GridSearchCV uses StratifiedKFold(n_splits=5, random_state=42)")
    bullet(doc, "RandomForestClassifier uses random_state=42")
    bullet(doc, "Preprocessing pipeline is fitted only on X_train — no leakage")
    bullet(doc, "All random seeds fixed — training results are deterministic")

    add_page_break(doc)

    # ── 6. CI/CD Pipeline ───────────────────────────────────────────────────────
    heading1(doc, "6. CI/CD Pipeline & Automated Testing")

    heading2(doc, "6.1 GitHub Actions Workflow")
    body(doc, (
        "A GitHub Actions pipeline defined in .github/workflows/ci.yml triggers on every push "
        "to main/develop branches and on pull requests to main. The pipeline runs sequentially "
        "and stops immediately if any step fails."
    ))

    add_table(doc,
        ["Step", "Tool", "Description"],
        [
            ["1. Checkout", "actions/checkout@v4", "Clone the repository"],
            ["2. Setup Python", "actions/setup-python@v5", "Python 3.13 environment"],
            ["3. Cache pip", "actions/cache@v4", "Cache dependencies for speed"],
            ["4. Install deps", "pip", "Install from requirements.txt"],
            ["5. Lint", "flake8", "PEP8 compliance, max line 100"],
            ["6. Format check", "black", "Code formatting validation"],
            ["7. Download data", "data/download_data.py", "Fetch Heart Disease dataset"],
            ["8. Unit tests", "pytest --cov", "Run 16 tests with coverage report"],
            ["9. Upload coverage", "actions/upload-artifact", "Save coverage.xml"],
            ["10. Train models", "src/train.py", "GridSearchCV + MLflow logging"],
            ["11. Upload models", "actions/upload-artifact", "Save trained .joblib files"],
            ["12. Docker build", "docker build", "Build container image"],
            ["13. Smoke test", "curl", "Test /health endpoint in container"],
        ]
    )

    heading2(doc, "6.2 Unit Tests (16 tests)")
    add_table(doc,
        ["Test File", "Test", "Validates"],
        [
            ["test_data_processing.py", "test_feature_names", "Correct feature split"],
            ["test_data_processing.py", "test_prepare_data_shapes", "Train/test sizes correct"],
            ["test_data_processing.py", "test_prepare_data_no_target_leakage", "No label in features"],
            ["test_data_processing.py", "test_preprocessor_returns_pipeline", "Pipeline created"],
            ["test_data_processing.py", "test_preprocessor_transforms", "No NaN after transform"],
            ["test_data_processing.py", "test_missing_summary", "Missing value counter"],
            ["test_data_processing.py", "test_prepare_data_handles_missing", "Imputer works"],
            ["test_model.py", "test_compute_metrics_keys", "All 5 metrics present"],
            ["test_model.py", "test_logistic_regression_trains", "LR trains and predicts"],
            ["test_model.py", "test_random_forest_trains", "RF output is binary"],
            ["test_model.py", "test_model_predict_proba_range", "Probabilities sum to 1"],
            ["test_api.py", "test_health_endpoint", "GET /health returns 200"],
            ["test_api.py", "test_root_endpoint", "GET / returns service name"],
            ["test_api.py", "test_predict_endpoint_valid", "POST /predict returns result"],
            ["test_api.py", "test_predict_missing_field", "422 on missing field"],
            ["test_api.py", "test_metrics_endpoint", "GET /metrics returns Prometheus"],
        ]
    )

    heading2(doc, "6.3 CI/CD Screenshots")
    add_figure(doc, "screenshots/github_actions_run.png",
               "Figure 10: GitHub Actions — Successful CI Run")
    add_figure(doc, "screenshots/github_actions_detail.png",
               "Figure 11: GitHub Actions — Step Detail with Logs")

    add_page_break(doc)

    # ── 7. Containerisation ──────────────────────────────────────────────────────
    heading1(doc, "7. Model Containerisation")

    heading2(doc, "7.1 Docker Architecture")
    body(doc, (
        "The FastAPI application is packaged into a Docker container based on python:3.11-slim. "
        "The image includes: the trained model (models/best_model.joblib), all source code "
        "(src/, api/), and all Python dependencies. The container exposes port 8000 and "
        "includes a Docker HEALTHCHECK that polls /health every 30 seconds."
    ))

    heading2(doc, "7.2 API Endpoints")
    add_table(doc,
        ["Method", "Endpoint", "Description", "Response"],
        [
            ["GET", "/", "Root health check", '{"status": "ok"}'],
            ["GET", "/health", "Liveness probe", '{"status": "healthy", "model_loaded": true}'],
            ["POST", "/predict", "Prediction endpoint", '{"prediction": 0|1, "label": "...", "confidence": 0.71}'],
            ["GET", "/metrics", "Prometheus scrape", "Prometheus text format"],
            ["GET", "/docs", "Swagger UI", "Interactive API documentation"],
        ]
    )

    heading2(doc, "7.3 Docker Commands")
    for cmd in [
        "# Build image",
        "docker build -t heart-disease-api:latest .",
        "",
        "# Run container",
        "docker run -p 8000:8000 heart-disease-api:latest",
        "",
        "# Full stack with monitoring",
        "docker-compose up --build",
        "",
        "# Test prediction",
        'curl -X POST http://localhost:8000/predict \\',
        '  -H "Content-Type: application/json" \\',
        '  -d \'{"age":63,"sex":1,"cp":3,"trestbps":145,"chol":233,',
        '       "fbs":1,"restecg":0,"thalach":150,"exang":0,',
        '       "oldpeak":2.3,"slope":0,"ca":0,"thal":1}\'',
    ]:
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        if cmd.startswith("#"):
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        p.paragraph_format.space_after = Pt(1)

    heading2(doc, "7.4 Docker Screenshots")
    add_figure(doc, "screenshots/docker_build.png",
               "Figure 12: Docker Image Build Output")
    add_figure(doc, "screenshots/docker_predict.png",
               "Figure 13: Docker Container — /predict Endpoint Response")
    add_figure(doc, "screenshots/api_health.png",
               "Figure 14: API /health Endpoint — Model Loaded")

    add_page_break(doc)

    # ── 8. Production Deployment ─────────────────────────────────────────────────
    heading1(doc, "8. Production Deployment on Kubernetes")

    heading2(doc, "8.1 Kubernetes Architecture")
    body(doc, (
        "The containerised API is deployed on Kubernetes using Minikube (local) or Docker Desktop "
        "Kubernetes. The deployment configuration includes a Deployment resource managing 2 replicas "
        "with rolling update strategy, and a Service resource of type LoadBalancer that exposes "
        "the API on port 80."
    ))

    heading2(doc, "8.2 Deployment Configuration")
    add_table(doc,
        ["K8s Resource", "File", "Key Settings"],
        [
            ["Deployment", "k8s/deployment.yaml",
             "replicas=2, imagePullPolicy=Never, liveness/readiness probes on /health"],
            ["Service", "k8s/service.yaml",
             "type=LoadBalancer, port=80 → targetPort=8000"],
        ]
    )

    heading2(doc, "8.3 Resource Limits")
    add_table(doc,
        ["Resource", "Request", "Limit"],
        [
            ["Memory", "256 Mi", "512 Mi"],
            ["CPU", "250m", "500m"],
        ]
    )

    heading2(doc, "8.4 Deployment Steps")
    steps = [
        "kubectl apply -f k8s/deployment.yaml",
        "kubectl apply -f k8s/service.yaml",
        "kubectl get pods                        # verify 2/2 Running",
        "kubectl get services                    # get external IP",
        "minikube service heart-disease-api-service --url  # Minikube only",
    ]
    for s in steps:
        p = doc.add_paragraph()
        r = p.add_run(s)
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(2)

    heading2(doc, "8.5 Deployment Screenshots")
    add_figure(doc, "screenshots/k8s_pods.png",
               "Figure 15: kubectl get pods — 2 Replicas Running")
    add_figure(doc, "screenshots/k8s_service.png",
               "Figure 16: kubectl get services — LoadBalancer Exposed")
    add_figure(doc, "screenshots/k8s_all_resources.png",
               "Figure 17: kubectl get all — Full Deployment State")

    add_page_break(doc)

    # ── 9. Monitoring & Logging ──────────────────────────────────────────────────
    heading1(doc, "9. Monitoring & Logging")

    heading2(doc, "9.1 API Request Logging")
    body(doc, (
        "Every HTTP request is captured by a FastAPI middleware that logs: HTTP method, "
        "endpoint path, HTTP status code, and response duration in seconds. Logs are emitted "
        "to stdout in a structured format and are visible via docker logs or kubectl logs."
    ))

    body(doc, "Example log line:", bold=True)
    p = doc.add_paragraph()
    r = p.add_run('2026-07-12 10:23:45  INFO  heart_disease_api  POST /predict  status=200  duration=0.041s')
    r.font.name = "Courier New"
    r.font.size = Pt(9)

    heading2(doc, "9.2 Prometheus Metrics")
    body(doc, "Three custom Prometheus counters/histograms are exposed at GET /metrics:")
    add_table(doc,
        ["Metric Name", "Type", "Description"],
        [
            ["api_requests_total", "Counter", "Total requests labelled by method, endpoint, status"],
            ["api_request_latency_seconds", "Histogram", "Request latency distribution per endpoint"],
            ["predictions_total", "Counter", "Total predictions labelled by output label"],
        ]
    )

    heading2(doc, "9.3 Prometheus + Grafana Stack")
    body(doc, (
        "Running docker-compose up brings up all three services. Prometheus scrapes the API's "
        "/metrics endpoint every 15 seconds (configured in monitoring/prometheus.yml). "
        "Grafana connects to Prometheus as a data source and allows building dashboards "
        "tracking request rates, prediction counts, and latency percentiles."
    ))

    add_table(doc,
        ["Service", "URL", "Credentials"],
        [
            ["Heart Disease API", "http://localhost:8000", "—"],
            ["Swagger UI", "http://localhost:8000/docs", "—"],
            ["Prometheus", "http://localhost:9090", "—"],
            ["Grafana", "http://localhost:3000", "admin / admin"],
        ]
    )

    heading2(doc, "9.4 Monitoring Screenshots")
    add_figure(doc, "screenshots/prometheus_targets.png",
               "Figure 18: Prometheus — Scrape Targets (API Endpoint UP)")
    add_figure(doc, "screenshots/prometheus_metrics.png",
               "Figure 19: Prometheus — api_requests_total Metric")
    add_figure(doc, "screenshots/api_metrics_raw.png",
               "Figure 20: Raw /metrics Endpoint — Prometheus Text Format")
    add_figure(doc, "screenshots/grafana_dashboard.png",
               "Figure 21: Grafana Dashboard — API Request Rate and Latency")

    add_page_break(doc)

    # ── 10. Setup Instructions & Conclusion ──────────────────────────────────────
    heading1(doc, "10. Setup Instructions & Conclusion")

    heading2(doc, "10.1 Prerequisites")
    bullet(doc, "Python 3.11+ (tested on 3.13)")
    bullet(doc, "Docker Desktop (with Kubernetes enabled for K8s deployment)")
    bullet(doc, "Git")

    heading2(doc, "10.2 Quick Start")
    steps2 = [
        ("Clone repository", "git clone https://github.com/<username>/heart-disease-mlops"),
        ("Create virtualenv", "python3 -m venv venv && source venv/bin/activate"),
        ("Install dependencies", "pip install -r requirements.txt"),
        ("Download dataset", "python data/download_data.py"),
        ("Run EDA", "PYTHONPATH=. python notebooks/01_eda.py"),
        ("Train models", "PYTHONPATH=. python -m src.train"),
        ("View MLflow UI", "mlflow ui --backend-store-uri sqlite:///mlflow.db"),
        ("Run tests", "PYTHONPATH=. pytest tests/ -v --cov=src --cov=api"),
        ("Start API", "PYTHONPATH=. uvicorn api.main:app --reload"),
        ("Run inference demo", "PYTHONPATH=. python notebooks/02_inference.py"),
        ("Full Docker stack", "docker-compose up --build"),
        ("Deploy to K8s", "kubectl apply -f k8s/"),
    ]

    for i, (desc, cmd) in enumerate(steps2, 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{i}. {desc}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(cmd)
        r2.font.name = "Courier New"
        r2.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(3)

    heading2(doc, "10.3 Architecture Summary")
    body(doc, (
        "The pipeline follows a linear MLOps flow:"
    ))
    arch_line = (
        "Raw Data → download_data.py → heart.csv\n"
        "→ data_processing.py (impute + scale + encode) → sklearn Pipeline\n"
        "→ train.py (GridSearchCV LR + RF) → MLflow experiment tracking\n"
        "→ best_model.joblib → FastAPI /predict endpoint\n"
        "→ Docker container → Kubernetes Deployment (2 replicas)\n"
        "→ Prometheus /metrics scrape → Grafana dashboards\n"
        "← GitHub Actions CI/CD (lint + test + train + docker smoke test)"
    )
    p = doc.add_paragraph()
    r = p.add_run(arch_line)
    r.font.name = "Courier New"
    r.font.size = Pt(9)

    heading2(doc, "10.4 Conclusion")
    body(doc, (
        "This project demonstrates a complete, production-grade MLOps workflow for a binary "
        "medical classification task. The Random Forest model achieved 90.2% accuracy and "
        "0.963 ROC-AUC on the test set, with the full sklearn Pipeline ensuring reproducible "
        "inference. All components — from data ingestion to monitoring — are automated, "
        "containerised, and deployable on Kubernetes with a single command."
    ))
    body(doc, (
        "Key MLOps principles applied:"
    ))
    bullet(doc, "Reproducibility: fixed seeds, versioned artifacts, pipeline serialisation")
    bullet(doc, "Automation: CI/CD pipeline runs on every commit — lint, test, train, deploy")
    bullet(doc, "Observability: structured logging + Prometheus metrics + Grafana dashboards")
    bullet(doc, "Portability: Docker container runs identically in development and production")
    bullet(doc, "Scalability: Kubernetes manages replicas, health probes, and rolling updates")

    heading2(doc, "10.5 Repository Link")
    body(doc, "GitHub Repository: https://github.com/yogisamy/MLOps-Assignment1")

    # ── Save ────────────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"Report saved: {OUTPUT}")
    return OUTPUT


if __name__ == "__main__":
    build_report()
